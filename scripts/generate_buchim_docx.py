#!/usr/bin/env python3
"""Generate 붙임2·붙임3 draft docx for drone-app submission."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUTPUT = "/Users/haru/drone-app/붙임2_붙임3_초안.docx"

APP_TITLE = "시뮬레이션으로 배우는 비행 — 드론 조종 코딩"
GRADE = "초등학교 실과 5학년"


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            set_cell_text(table.rows[r_idx + 1].cells[c_idx], val, size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 표지 정보
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("디지털교육연구대회\n연구보고서 붙임자료 (초안)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"\n{APP_TITLE}\n{GRADE}")
    r2.font.size = Pt(12)
    r2.font.name = "맑은 고딕"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_paragraph(
        "※ 본 문서는 붙임 2(멀티미디어 교육자료 목록)와 붙임 3(저작권 등 증빙자료) "
        "작성을 위한 초안입니다. 연구보고서 PDF 맨 뒤에 첨부하기 전에 "
        "출품자명·소속·증빙 캡처를 보완하세요."
    )

    # ── 핵심: Cursor AI 활용 개발 ──
    doc.add_paragraph()
    box_title = doc.add_paragraph()
    run_box = box_title.add_run("【 핵심 】 개발 도구 및 제작 방식")
    run_box.bold = True
    run_box.font.size = Pt(13)
    run_box.font.name = "맑은 고딕"
    run_box._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_paragraph(
        "본 소프트웨어의 모든 소스 코드(웹 애플리케이션, 3D 시뮬레이션, "
        "조종·코딩 모드, 효과음 합성, 오프라인 배포용 Windows·Android 패키지 설정 등)는 "
        "출품자가 AI 코딩 보조 도구 「Cursor AI」를 활용하여 기획·설계·작성·수정하였습니다."
    )
    doc.add_paragraph(
        "즉, 외부에서 구매하거나 내려받은 완성 프로그램을 그대로 사용한 것이 아니라, "
        "교육 목적에 맞게 출품자가 Cursor AI와 대화하며 요구사항을 전달하고, "
        "생성·검토·테스트를 반복하여 직접 완성한 결과물입니다."
    )

    cursor_items = [
        "개발 환경: Cursor AI + Next.js 16 + React 19 + Three.js",
        "제작 범위: 조종 모드, 코딩 모드, 3D 드론 시뮬레이션, UI/UX, "
        "Web Audio 효과음, PWA·오프라인(Windows run.bat, Android APK) 패키징",
        "저작 관계: 프로그램의 기획·교육 내용·최종 수정·검증은 출품자가 수행하였으며, "
        "코드 작성 과정에서 Cursor AI를 개발 도구로 활용함",
        "참고: Cursor AI는 개발 보조 도구이며, 본 소프트웨어에 포함된 "
        "외부 오픈소스 라이브러리(Three.js, React 등)의 저작권은 각 라이선스에 따름",
    ]
    for item in cursor_items:
        doc.add_paragraph(item, style="List Bullet")

    # ── 붙임 2 ──
    doc.add_page_break()
    h2 = doc.add_heading("붙임 2  멀티미디어 교육자료 목록", level=1)

    doc.add_paragraph(
        "■ media 폴더가 비어 있는 이유 (연구보고서 또는 본 붙임 서두에 기재 권장)"
    )
    doc.add_paragraph(
        "본 소프트웨어는 사진·동영상·음원 파일을 별도로 두지 않고, Three.js로 "
        "3D 화면(드론, 장애물, 바닥 격자 등)을, Web Audio API로 효과음(프로펠러 소리, "
        "성공·실패음)을 코드로 실시간 생성합니다. UI는 HTML·CSS로 구성합니다. "
        "따라서 제출 media 폴더에는 외부 미디어 파일이 없으며, 아래 목록은 "
        "프로그램에 실제로 사용·반영된 자료를 기준으로 작성하였습니다."
    )

    doc.add_paragraph("■ 멀티미디어 교육자료 목록")
    headers = ["번호", "자료형태", "파일명", "크기", "자료설명", "키워드", "비고(출처 등)"]

    rows_code = [
        ["1", "코드생성", "(해당 없음)", "-",
         "드론·프로펠러·장애물·목표지점·바닥 격자 등 3D 객체",
         "3D, 시뮬레이션",          "Three.js로 실시간 생성, 외부 이미지·모델 파일 미사용, Cursor AI로 코드 작성"],
        ["2", "코드생성", "(해당 없음)", "-",
         "프로펠러 소리·성공·실패 효과음",
         "음향, 효과음", "Web Audio API로 실시간 합성, Cursor AI로 코드 작성"],
        ["3", "코드생성", "(해당 없음)", "-",
         "버튼·조이스틱·UI 텍스트·배너",
         "UI, 화면구성", "HTML·CSS·Tailwind, Cursor AI로 코드 작성"],
    ]
    rows_files = [
        ["4", "그림", "icon.svg", "약 1KB",
         "앱 아이콘·PWA 아이콘", "아이콘", "출품자·Cursor AI 협업 제작"],
        ["5", "그림", "favicon.ico", "약 25KB",
         "브라우저 탭 아이콘", "아이콘", "Next.js 프로젝트 기본 아이콘"],
    ]
    rows_fonts = [
        ["6", "폰트", "Geist Sans", "-",
         "앱 전반 본문·UI 글꼴", "UI, 글꼴", "Google Fonts (next/font)"],
        ["7", "폰트", "Geist Mono", "-",
         "코드·숫자 표시용 글꼴", "UI, 글꼴", "Google Fonts (next/font)"],
    ]
    rows_modules = [
        ["8", "모듈", "three.js", "-",
         "3D 렌더링 엔진", "3D", "MIT License (npm)"],
        ["9", "모듈", "@react-three/fiber", "-",
         "React–Three.js 연동", "3D", "MIT License"],
        ["10", "모듈", "@react-three/drei", "-",
         "3D 보조 컴포넌트", "3D", "MIT License"],
        ["11", "모듈", "Next.js / React", "-",
         "웹 앱 프레임워크", "개발환경", "MIT License"],
    ]

    all_rows = rows_code + rows_files + rows_fonts + rows_modules
    add_table(doc, headers, all_rows, col_widths=[1.0, 1.6, 2.2, 1.2, 4.5, 2.0, 3.5])

    doc.add_paragraph()
    doc.add_paragraph(
        "※ public 폴더의 file.svg, vercel.svg 등 Next.js 기본 파일은 "
        "앱 화면에 사용하지 않으므로 목록에서 제외하였습니다."
    )

    # ── 붙임 3 ──
    doc.add_page_break()
    doc.add_heading("붙임 3  저작권 등 증빙자료", level=1)

    doc.add_paragraph(
        "아래 각 항목에 해당하는 화면 캡처·동의서·라이선스 페이지를 "
        "이어서 첨부하세요. (식물원 수상작 붙임 3 형식 참고)"
    )

    sections = [
        ("가. 직접 제작 자료 (Cursor AI 활용)", [
            "본 소프트웨어의 모든 소스 코드는 출품자가 Cursor AI를 개발 도구로 활용하여 "
            "기획·설계·작성·수정하였습니다. 3D 시뮬레이션 화면, 효과음, UI, "
            "icon.svg 및 오프라인 실행 패키지 설정까지 포함합니다.",
            "외부 이미지·동영상·음원 파일을 사용하지 않았으며, "
            "Three.js·Web Audio API·HTML/CSS로 실시간 생성합니다.",
            "[첨부] Cursor AI 개발 화면 캡처 (코드 작성·수정 과정, 권장)",
            "[첨부] icon.svg 원본 이미지 캡처 (선택)",
            "[첨부] 프로그램 주요 화면 캡처 1~2장 (선택)",
        ]),
        ("나. Google Fonts (Geist Sans, Geist Mono)", [
            "앱 전반 UI 글꼴로 Google Fonts의 Geist 계열을 사용하였습니다.",
            "[첨부] https://fonts.google.com/specimen/Geist 라이선스 화면 캡처",
            "[첨부] https://fonts.google.com/specimen/Geist+Mono 라이선스 화면 캡처",
        ]),
        ("다. 오픈소스 소프트웨어 (MIT License)", [
            "Three.js, React, Next.js, @react-three/fiber, @react-three/drei는 "
            "MIT License에 따라 사용하였으며, 상업적·교육적 이용이 허용됩니다.",
            "[첨부] https://github.com/mrdoob/three.js — LICENSE 화면 캡처",
            "[첨부] https://github.com/facebook/react — LICENSE 화면 캡처",
            "[첨부] https://github.com/vercel/next.js — license 화면 캡처",
        ]),
        ("라. 웹 배포 (수업용, 참고)", [
            "학생 수업에서는 설치 없이 바로 사용할 수 있도록 GitHub 저장소에 "
            "소스를 푸시하고 Vercel로 자동 배포하여 아래 웹 주소로 제공하였습니다.",
            "▶ https://drone-app-virid.vercel.app/",
            "[첨부] Vercel 배포 화면 또는 GitHub 저장소 화면 캡처 (선택)",
        ]),
        ("마. 제출 실행 파일 (참고)", [
            "오프라인 심사·제출 환경을 위해 다음 두 가지 실행 형태를 함께 제출하였습니다.",
            "(1) 윈도우 PC: run.bat + index.exe + out 폴더 (로컬 웹서버 방식, 완전 오프라인)",
            "(2) 안드로이드: 드론조종코딩.apk (WebView 래퍼, 완전 오프라인)",
        ]),
    ]

    for title_text, bullets in sections:
        doc.add_heading(title_text, level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("■ 첨부하지 않아도 되는 항목 (본 소프트웨어 미사용)")
    run.bold = True
    for item in [
        "YouTube·공공누리·국가생물종지식정보시스템 — 외부 사진·동영상 미사용",
        "미리캔버스·Unity Asset Store — 외부 그래픽 에셋 미사용",
        "ChatGPT API — 생성형 AI 기능 미사용",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "────────────────────────────────────────\n"
        "작성일: 2026년   월   일\n"
        "출품자:                    (서명)\n"
        "소속:                      \n"
        "────────────────────────────────────────"
    )

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
